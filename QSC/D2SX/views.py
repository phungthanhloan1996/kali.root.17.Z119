from decimal import Decimal
from .models import ProductItem
from collections import defaultdict
from .forms import ProductGroupForm, ProductItemForm, ProductForm
def create_product_group(request):
    if request.method == 'POST':
        form = ProductGroupForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('product_group_list')
    else:
        form = ProductGroupForm()
    return render(request, 'create_product_group.html', {'form': form})
def create_product_item(request):
    if request.method == 'POST':
        form = ProductItemForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('product_item_list')
    else:
        form = ProductItemForm()
    return render(request, 'create_product_item.html', {'form': form})
def create_product(request):
    if request.method == 'POST':
        form = ProductForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('product_list')
    else:
        form = ProductForm()
    return render(request, 'create_product.html', {'form': form})
def product_group_list(request):
    product_groups = ProductGroup.objects.all()
    return render(request, 'product_group_list.html', {'product_groups': product_groups})
def product_item_list(request):
    product_items = ProductItem.objects.all()
    for item in product_items:
        products = item.products.all()
        item.total_products = products.count()
        item.started_products = products.filter(start_date__isnull=False).count()
    return render(request, 'product_item_list.html', {'product_items': product_items})

from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.utils import timezone
from collections import defaultdict
from decimal import Decimal
import re

from .models import Product, ProductGroup, ProductItem


@login_required
def progress_list(request):
    today = timezone.localtime(timezone.now()).date()

    products = (
        Product.objects.select_related('item__group')
        .prefetch_related('progresses', 'item__steps', 'notes')
        .all()
    )
    product_groups = ProductGroup.objects.all().order_by('name')

    # --- TÍNH CHO TỪNG SẢN PHẨM ---
    for product in products:
        total_progress = Decimal(0)
        total_planned = Decimal(0)
        total_hours = Decimal(0)

        for step in product.item.steps.all():
            progress = product.progresses.filter(step=step).first()
            if progress:
                hours = Decimal(step.estimated_hours or 0)
                total_hours += hours
                total_progress += hours * Decimal(progress.progress_percent or 0)
                total_planned += hours * Decimal(progress.planned_progress_percent or 0)

        if total_hours > 0:
            product.total_progress = round(total_progress / total_hours, 2)
            product.planned_progress = round(total_planned / total_hours, 2)
        else:
            product.total_progress = Decimal(0)
            product.planned_progress = Decimal(0)

        unit_price = Decimal(product.item.unit_price or 0)
        product.total_value = round(unit_price * product.total_progress / 100, 2)
        product.planned_value = round(unit_price * product.planned_progress / 100, 2)
        product.total_value_when_completed = unit_price

    # --- TỔNG THEO NHÓM SẢN PHẨM ---
    for group in product_groups:
        group.products_in_group = [p for p in products if p.item.group == group]
        group.product_count = len(group.products_in_group)
        group.total_value_done = sum(p.total_value for p in group.products_in_group)
        group.total_planned_value = sum(p.planned_value for p in group.products_in_group)
        group.total_value_group = sum(p.total_value_when_completed for p in group.products_in_group)

    # --- CỘNG GỘP THEO MÃ NHÓM (có hoặc không có [prefix]) ---
    pattern = re.compile(r'^\[(.*?)\]')
    grouped_sums = defaultdict(lambda: {'done': Decimal(0), 'planned': Decimal(0), 'total': Decimal(0), 'members': []})

    for group in product_groups:
        match = pattern.match(group.name)
        prefix = match.group(1).strip().upper() if match else group.name.strip().upper()

        grouped_sums[prefix]['done'] += group.total_value_done
        grouped_sums[prefix]['planned'] += group.total_planned_value
        grouped_sums[prefix]['total'] += group.total_value_group
        grouped_sums[prefix]['members'].append(group.name)

        if 'name' not in grouped_sums[prefix]:
            clean_name = re.sub(r'^\[[^\]]+\]\s*', '', group.name).strip()
            grouped_sums[prefix]['name'] = clean_name or group.name

    # --- TỔNG CỘNG TOÀN BỘ ---
    total_done_all = float(sum(float(data['done'] or 0) for data in grouped_sums.values()))
    total_planned_all = float(sum(float(data['planned'] or 0) for data in grouped_sums.values()))
    total_value_all = float(sum(float(data['total'] or 0) for data in grouped_sums.values()))

    product_items = ProductItem.objects.all()
    product_count_per_item = defaultdict(int)
    for product in products:
        product_count_per_item[product.item.id] += 1

    context = {
        'today': today,
        'products': products,
        'product_groups': product_groups,
        'product_items': product_items,
        'product_count_per_item': product_count_per_item,
        'grouped_sums': grouped_sums,
        'grouped_sums': dict(grouped_sums),
        'total_done_all': total_done_all,
        'total_planned_all': total_planned_all,
        'total_value_all': total_value_all,
    }
    return render(request, 'progress_list.html', context)

# ngày 05/6/205 Phùng Thái Hà UPDATE
# views.py — update_progress (thay thế hoàn toàn phần hiện tại)
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.utils import timezone
from datetime import datetime

from .models import Product, ProductStepProgress, ProductNote, ProductDailyProgress

def update_progress(request, product_id):
    product = get_object_or_404(Product, id=product_id)
    steps = product.item.steps.order_by('order')

    # --- Lấy selected_date từ GET (nếu có) hoặc POST, default = today ---
    selected_date_str = request.GET.get('selected_date') or request.POST.get('selected_date')
    if selected_date_str:
        try:
            selected_date = datetime.strptime(selected_date_str, "%Y-%m-%d").date()
        except Exception:
            selected_date = timezone.localtime(timezone.now()).date()
    else:
        selected_date = timezone.localtime(timezone.now()).date()

    if request.method == 'POST':
        # Lặp qua các step, cập nhật ProductStepProgress (cũ của bạn)
        for step in steps:
            progress_value = request.POST.get(f'progress_{step.id}')
            planned_value = request.POST.get(f'planned_{step.id}')
            note_actual = request.POST.get(f'note_actual_{step.id}')
            note_planned = request.POST.get(f'note_planned_{step.id}')

            progress_obj = product.progresses.filter(step=step).first()
            if progress_obj:
                # giữ quyền control như cũ
                if not request.user.is_leader:
                    progress_value = progress_obj.progress_percent
                    note_actual = progress_obj.note_actual
                if not request.user.is_director:
                    planned_value = progress_obj.planned_progress_percent
                    note_planned = progress_obj.note_planned

                if request.user.is_leader and progress_value not in (None, ''):
                    try:
                        progress_obj.progress_percent = float(progress_value)
                    except ValueError:
                        pass
                    progress_obj.note_actual = note_actual or ''
                if request.user.is_director and planned_value not in (None, ''):
                    try:
                        progress_obj.planned_progress_percent = float(planned_value)
                    except ValueError:
                        pass
                    progress_obj.note_planned = note_planned or ''

                # cập nhật completed_date nếu cần (save method của model vẫn chạy)
                progress_obj.save()
            else:
                # tạo mới ProductStepProgress nếu chưa có
                pdata = {
                    'product': product,
                    'step': step,
                    'progress_percent': 0,
                    'planned_progress_percent': 0,
                    'note_actual': '',
                    'note_planned': '',
                }
                if request.user.is_leader:
                    try:
                        pdata['progress_percent'] = float(progress_value or 0)
                    except ValueError:
                        pass
                    pdata['note_actual'] = note_actual or ''
                if request.user.is_director:
                    try:
                        pdata['planned_progress_percent'] = float(planned_value or 0)
                    except ValueError:
                        pass
                    pdata['note_planned'] = note_planned or ''

                progress_obj = ProductStepProgress.objects.create(**pdata)

        # Lưu product note (giữ như cũ)
        note = request.POST.get('product_note')
        if note:
            ProductNote.objects.create(product=product, content=note, created_by=request.user)

        # --- CẬP NHẬT ProductDailyProgress cho selected_date ---
        # Tính tổng thực tế theo bước: trung bình có trọng số theo estimated_hours (dùng hàm total_progress nếu bạn có)
        try:
            # Nếu model Product.total_progress() dùng ProductStepProgress (không theo ngày),
            # thì nó phản ánh "latest step progress" — chấp nhận được nếu bạn muốn.
            aggregate_real = product.total_progress()  # trả về số 0..100
        except Exception:
            # fallback: trung bình đơn giản
            steps_qs = product.item.steps.all()
            if steps_qs:
                vals = []
                for s in steps_qs:
                    p = product.progresses.filter(step=s).first()
                    vals.append(float(p.progress_percent) if p else 0)
                aggregate_real = sum(vals) / len(vals) if vals else 0
            else:
                aggregate_real = 0

        # tính aggregate_plan (dùng planned_progress_percent từ step progress nếu có)
        try:
            planned_vals = []
            for s in product.item.steps.all():
                p = product.progresses.filter(step=s).first()
                planned_vals.append(float(p.planned_progress_percent) if p else 0)
            aggregate_plan = sum(planned_vals) / len(planned_vals) if planned_vals else 0
        except Exception:
            aggregate_plan = 0

        # get or create ProductDailyProgress cho ngày selected_date
        dp, created = ProductDailyProgress.objects.get_or_create(product=product, date=selected_date, defaults={
            'progress_percent': aggregate_real,
            'planned_progress_percent': aggregate_plan,
            'approved_factory': False,
            'approved_department': False,
            'approved_qc': False,
            'handed_over': False,
        })
        # nếu tồn tại, cập nhật các giá trị (chúng ta không ghi đè flags nếu user không gửi, nhưng bạn có thể)
        dp.progress_percent = aggregate_real
        dp.planned_progress_percent = aggregate_plan
        # (nếu muốn, lấy checkbox từ form để cập nhật approved_* — nhưng form hiện tại không có, nên giữ nguyên)
        dp.save()

        messages.success(request, 'Lưu tiến độ theo ngày thành công.')
        # redirect về cùng trang và giữ selected_date trong querystring
        return redirect(f"{request.path}?selected_date={selected_date.isoformat()}")

    # Khi GET: truyền dữ liệu để hiển thị form
    # Lấy progress hiện tại để hiển thị giá trị mặc định trong form (không theo ngày)
    progresses = {p.step.id: p for p in product.progresses.all()}
    latest_note = ProductNote.objects.filter(product=product).order_by('-created_at').first()

    return render(request, 'update_progress.html', {
        'product': product,
        'steps': steps,
        'progresses': progresses,
        'latest_note': latest_note,
        'selected_date': selected_date,
        'today': timezone.localtime(timezone.now()).date(),
    })



# NGÀY 01/6/2025 PHÙNG THÁI HÀ UPDATE
from django.shortcuts import render, get_object_or_404, redirect
from .models import Product, ProductStepProgress
def update_product_progress(request, product_id):
    product = get_object_or_404(Product, id=product_id)
    steps = product.item.steps.all()
    progresses = {progress.step.id: progress.progress_percent for progress in product.progresses.all()}
    if request.method == 'POST':
        for step in steps:
            progress_value = request.POST.get(f'progress_{step.id}')
            if progress_value:
                try:
                    progress_value = float(progress_value)
                    ProductStepProgress.objects.update_or_create(
                        product=product,
                        step=step,
                        defaults={'progress_percent': progress_value}
                    )
                except ValueError:
                    pass
        return redirect('progress_list')
    return render(request, 'update_progress.html', {
        'product': product,
        'steps': steps,
        'progresses': progresses
    })


from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.shortcuts import render, redirect
from .models import Product, OrderForm, OrderItem, Z119group
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
from django.http import HttpResponse

@login_required
def order_form(request):
    user = request.user
    product_dh = Product.objects.all()
    z119_groups = Z119group.objects.all()

    # Lấy tên nhóm của user (đơn vị đặt)
    z119group_name = getattr(user, 'z119group', None)
    if z119group_name:
        z119group_name = z119group_name.name
    else:
        z119group_name = ''

    if request.method == "POST":
        submit_type = request.POST.get("submit_type", "save")  # mặc định là lưu phiếu

        # Lấy dữ liệu chung
        executing_unit_name = request.POST.get("executing_unit")
        norm_unit_name = request.POST.get("norm_unit")
        order_date = request.POST.get("order_date")

        # Validate đơn vị thực hiện
        try:
            executing_unit = Z119group.objects.get(name=executing_unit_name)
        except Z119group.DoesNotExist:
            messages.error(request, "Đơn vị thực hiện không hợp lệ.")
            return redirect("order_form")

        # Validate đơn vị định mức
        try:
            norm_unit = Z119group.objects.get(name=norm_unit_name)
        except Z119group.DoesNotExist:
            messages.error(request, "Đơn vị định mức không hợp lệ.")
            return redirect("order_form")

        # Thu thập các dòng hàng hóa
        items = []
        MAX_ROWS = 30
        for row_num in range(1, MAX_ROWS + 1):
            item_name = request.POST.get(f"item_name_{row_num}")
            quantity_str = request.POST.get(f"quantity_{row_num}")
            product_code = request.POST.get(f"product_code_{row_num}")
            due_date = request.POST.get(f"due_date_{row_num}")
            note = request.POST.get(f"note_{row_num}")

            if not any([item_name, quantity_str, product_code, due_date, note]):
                break

            if not quantity_str or not due_date:
                continue

            try:
                quantity = int(quantity_str)
                if quantity <= 0:
                    continue
            except ValueError:
                continue

            product_obj = None
            product_name = ""
            if product_code:
                try:
                    product_obj = Product.objects.get(id=product_code)
                    product_name = product_obj.name
                except (ValueError, Product.DoesNotExist):
                    pass

            items.append({
                'tt': str(len(items) + 1),
                'ten_hang': item_name or product_name,
                'sl': quantity,
                'so_viec': product_code,
                'ngay_xong': due_date,
                'ghi_chu': note or ""
            })

        # PHÂN BIỆT HÀNH ĐỘNG
        if submit_type == "export_word":
            # XUẤT RA WORD
            doc = Document()

            def set_run_font(run, size=12, bold=False):
                run.font.name = 'Times New Roman'
                run.font.size = Pt(size)
                run.bold = bold
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            # Tiêu đề
            p = doc.add_paragraph("PHIẾU ĐẶT HÀNG")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.runs[0], size=16, bold=True)

            doc.add_paragraph()  # khoảng trống

            # Thông tin đầu
            p = doc.add_paragraph()
            run = p.add_run("Đơn vị đặt: ")
            set_run_font(run, bold=True)
            p.add_run(z119group_name)

            p = doc.add_paragraph()
            run = p.add_run("Đơn vị thực hiện: ")
            set_run_font(run, bold=True)
            p.add_run(executing_unit_name)

            p = doc.add_paragraph()
            run = p.add_run("Đơn vị định mức: ")
            set_run_font(run, bold=True)
            p.add_run(norm_unit_name)

            p = doc.add_paragraph()
            run = p.add_run("Ngày đặt: ")
            set_run_font(run, bold=True)
            p.add_run(order_date)

            doc.add_paragraph()

            # Bảng chi tiết
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            headers = ['TT', 'TÊN HÀNG', 'SL', 'SỐ VIỆC', 'NGÀY XONG', 'GHI CHÚ']
            hdr_cells = table.rows[0].cells
            for i, text in enumerate(headers):
                run = hdr_cells[i].paragraphs[0].add_run(text)
                set_run_font(run, bold=True)
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for item in items:
                row = table.add_row().cells
                row[0].text = item['tt']
                row[1].text = item['ten_hang']
                row[2].text = str(item['sl'])
                row[3].text = item['so_viec']
                row[4].text = item['ngay_xong']
                row[5].text = item['ghi_chu']

                for cell in row:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell.paragraphs[0].runs:
                        set_run_font(run)

            # Điều chỉnh độ rộng cột (có thể chỉnh lại tùy ý để giống ảnh hơn)
            column_widths_inches = [0.4, 3.0, 0.8, 1.5, 1.2, 2.0]
            for i, width in enumerate(column_widths_inches):
                table.columns[i].width = Inches(width)

            # Tạo response Word
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            response = HttpResponse(
                content=buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = 'attachment; filename="Phieu_Dat_Hang.docx"'
            return response

        else:
            # LƯU PHIẾU (hành động mặc định)
            order_form_obj = OrderForm.objects.create(
                created_by=user,
                ordering_unit=user.z119group,
                executing_unit=executing_unit,
                norm_unit=norm_unit,
                order_date=order_date,
                status='submitted'
            )

            for item in items:
                product_obj = None
                if item['so_viec']:
                    try:
                        product_obj = Product.objects.get(id=item['so_viec'])
                    except:
                        pass

                OrderItem.objects.create(
                    order_form=order_form_obj,
                    item_name=item['ten_hang'],
                    quantity=item['sl'],
                    product=product_obj,
                    due_date=item['ngay_xong'],
                    note=item['ghi_chu']
                )

            messages.success(request, "Đã gửi phiếu đặt hàng.")
            return redirect("order_list")

    # GET request
    context = {
        "user_group": z119group_name,
        "z119_groups": z119_groups,
        "product_dh": product_dh
    }
    return render(request, "PXdathang.html", context)

from django.db.models import Count
from django.contrib.auth.decorators import login_required

@login_required
def order_list(request):
    user = request.user
    if user.z119group and user.z119group.name == "Ban Giám đốc":
        orders = OrderForm.objects.filter(status='submitted').prefetch_related('items__product').order_by('-order_date')
    else:
        orders = OrderForm.objects.filter(created_by=user).prefetch_related('items__product').order_by('-order_date')

    return render(request, 'order_list.html', {'orders': orders})


@login_required
def order_detail(request, pk):
    order = get_object_or_404(OrderForm, pk=pk)

    user = request.user
    if not (order.created_by == user or (user.z119group and user.z119group.name == "Phòng Kế hoạch")):
        messages.error(request, "Bạn không có quyền xem phiếu này.")
        return redirect('order_list')

    items = order.items.all()
    return render(request, 'order_detail.html', {
        'order': order,
        'items': items,
    })

@login_required
def order_norm_list(request):
    orders = OrderForm.objects.filter(
        status='forwarded',
        norm_unit=request.user.z119group
    ).order_by('-order_date')
    return render(request, 'order_norm_list.html', {'orders': orders})


def is_director(user):
    return user.is_authenticated and user.is_director

from django.contrib.auth.decorators import user_passes_test
from django.utils import timezone
from django.shortcuts import get_object_or_404, redirect, render
from django.contrib import messages


@login_required
@user_passes_test(is_director)
def order_review_list(request):
    orders = OrderForm.objects.filter(status='submitted') \
        .select_related('ordering_unit', 'executing_unit', 'norm_unit') \
        .order_by('-order_date')

    return render(request, 'order_review_list.html', {'orders': orders})


@login_required
@user_passes_test(is_director)
def order_review_detail(request, pk):
    order = get_object_or_404(OrderForm, pk=pk)

    if request.method == 'POST':
        action = request.POST.get('action')
        if action == 'approve':
            order.status = 'forwarded'
            order.reviewed_by = request.user
            order.reviewed_at = timezone.now()
            order.save()
            messages.success(request, "Phiếu đã được duyệt.")
        elif action == 'reject':
            order.status = 'rejected'
            order.reviewed_by = request.user
            order.reviewed_at = timezone.now()
            order.save()
            messages.success(request, "Phiếu đã bị từ chối.")

        return redirect('order_review_list')

    items = order.items.all()
    return render(request, 'order_review_detail.html', {'order': order, 'items': items})


from django.contrib.auth.decorators import login_required, user_passes_test
from django.shortcuts import render, get_object_or_404
from .models import OrderForm

def is_leader(user):
    return user.is_authenticated and user.is_leader
    
from django.db.models import Count, Q

from django.db.models import Count, Q
@login_required
@user_passes_test(is_leader)
def order_execution_list(request):
    orders = (
        OrderForm.objects.filter(
            status__in=['forwarded', 'completed'],
            executing_unit=request.user.z119group
        )
        .annotate(
            total_items=Count('items'),
            completed_items=Count('items', filter=Q(items__is_completed=True))
        )
        .order_by('-order_date')
    )
    return render(request, 'order_execution_list.html', {'orders': orders})

from django.utils import timezone

from django.contrib.auth.decorators import login_required, user_passes_test
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.http import HttpResponse
from django.utils import timezone
from .models import OrderForm, OrderItem
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
import io

# Giả sử hàm is_leader đã được định nghĩa ở đâu đó
def is_leader(user):
    return user.is_authenticated and user.z119group and user.z119group.is_leader  # ví dụ

@login_required
#@user_passes_test(is_leader)
def order_execution_detail(request, pk):
    order_form = get_object_or_404(
        OrderForm,
        pk=pk,
        executing_unit=request.user.z119group
    )
    items = order_form.items.all()

    if request.method == "POST":
        submit_type = request.POST.get("submit_type", "mark_complete")

        if submit_type == "export_word":
            # ── XUẤT WORD ────────────────────────────────────────────────────────────────
            doc = Document()

            def set_run_font(run, size=12, bold=False):
                run.font.name = 'Times New Roman'
                run.font.size = Pt(size)
                run.bold = bold
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            # Tiêu đề
            p = doc.add_paragraph(f"CHI TIẾT PHIẾU ĐẶT HÀNG (THỰC HIỆN) #{order_form.id}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.runs[0], size=16, bold=True)

            doc.add_paragraph()

            # Thông tin phiếu
            p = doc.add_paragraph()
            run = p.add_run("Đơn vị đặt: ")
            set_run_font(run, bold=True)
            p.add_run(order_form.ordering_unit.name if order_form.ordering_unit else "N/A")

            p = doc.add_paragraph()
            run = p.add_run("Đơn vị định mức: ")
            set_run_font(run, bold=True)
            p.add_run(order_form.norm_unit.name if order_form.norm_unit else "N/A")

            p = doc.add_paragraph()
            run = p.add_run("Đơn vị thực hiện: ")
            set_run_font(run, bold=True)
            p.add_run(order_form.executing_unit.name if order_form.executing_unit else "N/A")

            p = doc.add_paragraph()
            run = p.add_run("Ngày đặt: ")
            set_run_font(run, bold=True)
            p.add_run(str(order_form.order_date) if order_form.order_date else "N/A")

            p = doc.add_paragraph()
            run = p.add_run("Trạng thái phiếu: ")
            set_run_font(run, bold=True)
            p.add_run(order_form.get_status_display())

            doc.add_paragraph()

            # Bảng danh sách hàng
            table = doc.add_table(rows=1, cols=9)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            headers = [
                'STT', 'Tên hàng', 'Số lượng', 'Số việc / Sản phẩm',
                'Ngày phải xong', 'Ngày hoàn thành', 'Ghi chú', 'Trạng thái', 'Ghi chú thêm'
            ]
            hdr_cells = table.rows[0].cells
            for i, text in enumerate(headers):
                run = hdr_cells[i].paragraphs[0].add_run(text)
                set_run_font(run, bold=True)
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for idx, item in enumerate(items, start=1):
                row = table.add_row().cells
                row[0].text = str(idx)
                row[1].text = item.item_name or ""
                row[2].text = str(item.quantity)
                row[3].text = item.product.name if item.product else (item.item_name or "-")
                row[4].text = str(item.due_date) if item.due_date else "-"
                row[5].text = str(item.completed_date) if item.completed_date else "-"
                row[6].text = item.note or ""
                row[7].text = "Hoàn thành" if item.is_completed else "Chưa hoàn thành"
                row[8].text = ""  # cột trống để giống form gốc nếu cần

                for cell in row:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell.paragraphs[0].runs:
                        set_run_font(run)

            # Độ rộng cột (có thể điều chỉnh theo ý)
            column_widths_inches = [0.4, 2.5, 0.9, 2.0, 1.2, 1.2, 1.8, 1.2, 0.8]
            for i, width in enumerate(column_widths_inches):
                table.columns[i].width = Inches(width)

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            response = HttpResponse(
                content=buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="Phieu_Thuc_Hien_{order_form.id}.docx"'
            return response

        else:
            # Xử lý đánh dấu hoàn thành (logic cũ)
            item_id = request.POST.get("item_id")
            try:
                item = order_form.items.get(pk=item_id)
                item.is_completed = True
                item.completed_date = timezone.now().date()
                item.save()
                messages.success(request, f"Mặt hàng '{item.item_name}' đã được đánh dấu hoàn thành.")

                if all(i.is_completed for i in order_form.items.all()):
                    order_form.status = "completed"
                    order_form.save()
                    messages.success(request, "Toàn bộ phiếu đã hoàn thành.")

            except OrderItem.DoesNotExist:
                messages.error(request, "Không tìm thấy mặt hàng này.")

            return redirect("order_execution_detail", pk=order_form.pk)

    # GET request
    return render(request, "order_execution_detail.html", {
        "order_form": order_form,
        "items": items,
    })

from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from .models import OrderForm

def is_norm_unit(user):
    return user.is_authenticated and user.z119group is not None

@login_required
@user_passes_test(is_norm_unit)
def order_norm_detail(request, pk):
    order = get_object_or_404(OrderForm, pk=pk, norm_unit=request.user.z119group)
    items = order.items.all()

    if request.method == 'POST' and request.POST.get('action') == 'export_excel':
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
        from openpyxl.worksheet.page import PageMargins
        from datetime import date

        wb = Workbook()
        ws = wb.active
        ws.title = f"DM-{order.id}"

        # Fonts
        f14b = Font(name='Times New Roman', size=14, bold=True)
        f14bu = Font(name='Times New Roman', size=14, bold=True, underline='single')
        f12b = Font(name='Times New Roman', size=12, bold=True)   # tăng size nhẹ cho chữ ký
        f11b = Font(name='Times New Roman', size=11, bold=True)
        f11 = Font(name='Times New Roman', size=11)
        f11i = Font(name='Times New Roman', size=11, italic=True)

        # Alignments
        ac = Alignment(horizontal='center', vertical='center', wrap_text=True)
        al = Alignment(horizontal='left', vertical='center', wrap_text=True)
        ar = Alignment(horizontal='right', vertical='center', wrap_text=True)

        # Border
        def tb():
            s = Side(style='thin')
            return Border(left=s, right=s, top=s, bottom=s)

        header_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")

        # Độ rộng cột - điều chỉnh cho đẹp, cân đối
        ws.column_dimensions['A'].width = 6      # TT
        ws.column_dimensions['B'].width = 45     # Nội dung công việc - rộng nhất
        ws.column_dimensions['C'].width = 14     # Số lượng
        ws.column_dimensions['D'].width = 14     # Thời gian (giờ)
        ws.column_dimensions['E'].width = 20     # Ghi chú - rộng hơn tí

        # Row 1
        ws.merge_cells('A1:E1')
        ws['A1'] = "QC PHÒNG KHÔNG - KHÔNG QUÂN"
        ws['A1'].font = f14b
        ws['A1'].alignment = al
        ws.row_dimensions[1].height = 24

        # Row 2
        ws.merge_cells('A2:E2')
        ws['A2'] = "NHÀ MÁY Z119"
        ws['A2'].font = f14bu
        ws['A2'].alignment = al
        ws.row_dimensions[2].height = 24
        ws.row_dimensions[3].height = 8

        # Row 4: Tiêu đề
        ws.merge_cells('A4:E4')
        ws['A4'] = "ĐỊNH MỨC THỜI GIAN"
        ws['A4'].font = f14b
        ws['A4'].alignment = ac
        ws.row_dimensions[4].height = 28
        ws.row_dimensions[5].height = 8

        # Row 6: Sản phẩm + Số việc
        first_item = items.first()
        product_name = first_item.item_name if first_item else ""
        so_viec = first_item.product.name if first_item and first_item.product else ""
        ws.merge_cells('A6:C6')
        ws['A6'] = f"Sản phẩm: {product_name}"
        ws['A6'].font = f11
        ws['A6'].alignment = al
        ws.merge_cells('D6:E6')
        ws['D6'] = f"Số việc: {so_viec}"
        ws['D6'].font = f11
        ws['D6'].alignment = ar   # căn phải cho đẹp
        ws.row_dimensions[6].height = 22

        # Row 7: Đơn vị
        ws.merge_cells('A7:C7')
        ws['A7'] = f"Đơn vị thực hiện: {order.executing_unit.name if order.executing_unit else ''}"
        ws['A7'].font = f11
        ws['A7'].alignment = al
        ws.merge_cells('D7:E7')
        ws['D7'] = f"Đơn vị đặt hàng: {order.ordering_unit.name if order.ordering_unit else ''}"
        ws['D7'].font = f11
        ws['D7'].alignment = ar
        ws.row_dimensions[7].height = 22
        ws.row_dimensions[8].height = 10

        # Header bảng row 9
        headers = [
            ('A', 'TT'),
            ('B', 'Nội dung công việc'),
            ('C', 'Số lượng'),
            ('D', 'Thời gian\n(giờ)'),
            ('E', 'Ghi chú')
        ]
        for col, val in headers:
            c = ws[f'{col}9']
            c.value = val
            c.font = f11b
            c.alignment = ac
            c.border = tb()
            c.fill = header_fill
        ws.row_dimensions[9].height = 40   # cao hơn tí cho header đẹp

        # Dữ liệu items
        current_row = 10
        for idx, item in enumerate(items, 1):
            row = current_row
            ws.cell(row=row, column=1).value = str(idx)
            ws.cell(row=row, column=2).value = item.item_name or ""

            # Số lượng
            qty = ""
            if hasattr(item, 'quantity'):
                qty = item.quantity
            elif hasattr(item, 'so_luong'):
                qty = item.so_luong
            elif hasattr(item, 'qty'):
                qty = item.qty
            ws.cell(row=row, column=3).value = qty

            ws.cell(row=row, column=4).value = ""  # Thời gian để trống
            ws.cell(row=row, column=5).value = item.note or ""

            # Alignment
            ws.cell(row=row, column=1).alignment = ac
            ws.cell(row=row, column=2).alignment = al
            ws.cell(row=row, column=3).alignment = ac
            ws.cell(row=row, column=4).alignment = ac
            ws.cell(row=row, column=5).alignment = al

            for col in range(1, 6):
                ws.cell(row=row, column=col).font = f11
                ws.cell(row=row, column=col).border = tb()

            ws.row_dimensions[row].height = 38   # cao hơn tí cho dễ đọc

            current_row += 1

        # Để khoảng trống trước chữ ký
        current_row += 2
        ws.row_dimensions[current_row].height = 12
        current_row += 1

        # Ngày tháng - căn phải
        ws.merge_cells(f'A{current_row}:E{current_row}')
        today = date.today()
        ws[f'A{current_row}'] = f"Hà Nội, ngày {today.day} tháng {today.month:02d} năm {today.year}"
        ws[f'A{current_row}'].font = f11i
        ws[f'A{current_row}'].alignment = ar
        ws.row_dimensions[current_row].height = 24
        current_row += 3   # tăng khoảng cách

        # Chữ ký - thiết kế lại đẹp hơn
        sign_row = current_row

        # Cột A: KT. GIÁM ĐỐC + Phó
        ws.merge_cells(f'A{sign_row}:A{sign_row+1}')
        ws[f'A{sign_row}'] = "KT. GIÁM ĐỐC\nPHÓ GIÁM ĐỐC"
        ws[f'A{sign_row}'].font = f12b
        ws[f'A{sign_row}'].alignment = ac

        # Cột B+C: KT. TP KỸ THUẬT + Phó
        ws.merge_cells(f'B{sign_row}:C{sign_row}')
        ws.merge_cells(f'B{sign_row+1}:C{sign_row+1}')
        ws[f'B{sign_row}'] = "KT. TRƯỞNG PHÒNG KỸ THUẬT"
        ws[f'B{sign_row+1}'] = "PHÓ TRƯỞNG PHÒNG KỸ THUẬT"
        ws[f'B{sign_row}'].font = f12b
        ws[f'B{sign_row+1}'].font = f12b
        ws[f'B{sign_row}'].alignment = ac
        ws[f'B{sign_row+1}'].alignment = ac

        # Cột D+E: NGƯỜI LẬP (gộp rộng hơn)
        ws.merge_cells(f'D{sign_row}:E{sign_row+1}')
        ws[f'D{sign_row}'] = "NGƯỜI LẬP BIỂU"
        ws[f'D{sign_row}'].font = f12b
        ws[f'D{sign_row}'].alignment = ac

        # Tên người ký (dòng dưới)
        name_row = sign_row + 3
        ws.merge_cells(f'A{name_row}:A{name_row}')
        ws[f'A{name_row}'] = "Thượng tá Nguyễn Bá Quang"
        ws[f'A{name_row}'].font = f11b
        ws[f'A{name_row}'].alignment = ac

        ws.merge_cells(f'B{name_row}:C{name_row}')
        ws[f'B{name_row}'] = "Trung tá Đặng Việt Hưng"
        ws[f'B{name_row}'].font = f11b
        ws[f'B{name_row}'].alignment = ac

        ws.merge_cells(f'D{name_row}:E{name_row}')
        ws[f'D{name_row}'] = ""   # để trống cho người dùng tự điền tên
        ws[f'D{name_row}'].font = f11b
        ws[f'D{name_row}'].alignment = ac

        # Border cho phần chữ ký (tùy chọn - thêm border nhẹ)
        for r in range(sign_row, sign_row + 4):
            for col in range(1, 6):
                if ws.cell(row=r, column=col).value:  # chỉ border ô có chữ
                    ws.cell(row=r, column=col).border = tb()

        # Cài đặt in
        ws.page_setup.orientation = 'portrait'
        ws.page_setup.paperSize = 9  # A4
        ws.page_margins = PageMargins(left=0.6, right=0.6, top=0.8, bottom=0.8, header=0.3, footer=0.3)
        ws.print_area = f'A1:E{name_row + 2}'
        ws.sheet_view.showGridLines = False

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="Dinh_muc_thoi_gian_{order.id}.xlsx"'
        wb.save(response)
        return response

    return render(request, 'order_norm_detail.html', {'order': order, 'items': items})



def product_list_view(request):
    # Lấy ra danh sách sản phẩm có xuất hiện trong phiếu đặt hàng
    products = Product.objects.filter(orderitem__product__isnull=False).distinct()
    return render(request, 'products_with_orders.html', {'products': products})


from django.shortcuts import render, get_object_or_404
from .models import Product, OrderItem, OrderForm

def order_items_by_product(request, product_id):
    # Lấy sản phẩm cần xem
    product = get_object_or_404(Product, id=product_id)

    # Lấy tất cả OrderItem gắn với sản phẩm này
    order_items = (
        OrderItem.objects
        .filter(product=product)
        .select_related('order_form')
    )

    # Lấy danh sách phiếu liên quan
    order_forms = (
        OrderForm.objects
        .filter(items__product=product)
        .distinct()
        .select_related('executing_unit', 'norm_unit', 'ordering_unit')
        .prefetch_related('items')
    )

    context = {
        "product": product,
        "order_items": order_items,
        "order_forms": order_forms,
    }
    return render(request, "order_items_by_product.html", context)

from .models import ProductGroup, Product
def product_list_by_group(request):
    groups = ProductGroup.objects.all()
    group_products = {}
    for group in groups:
        products = Product.objects.filter(item__group=group).select_related('item')
        group_products[group] = products
    return render(request, 'product_list_by_group.html', {
        'group_products': group_products,
    })


def order_form_detail(request, order_form_id):
    order_form = get_object_or_404(OrderForm, id=order_form_id)
    items = order_form.items.all()
    return render(request, 'order_form_detail.html', {
        'order_form': order_form,
        'items': items
    })


from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.utils import timezone
from datetime import timedelta
from django.utils.timezone import localtime
from django.http import JsonResponse

from .models import (
    ProductGroup, Product, ProductDailyProgress,
    ProductStepProgress, ProductNote
)
from DKQSZ119.models import AttendanceRecord, WorkerInfo

@login_required
def dashboard_sum(request):
    today = timezone.localtime(timezone.now()).date()

    # --- Nếu POST là gửi note AJAX ---
    if request.method == 'POST':
        product_id = request.POST.get('product_id')
        content = request.POST.get('content', '').strip()
        if product_id and content:
            try:
                product = Product.objects.get(id=product_id)
                note = ProductNote.objects.create(
                    product=product,
                    content=content,
                    created_by=request.user
                )
                # Trả JSON nếu là AJAX
                if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                    return JsonResponse({
                        'success': True,
                        'content': note.content,
                        'author': note.created_by.get_full_name() or note.created_by.username,
                        'created_at': localtime(note.created_at).strftime("%d/%m/%Y %H:%M")
                    })
            except Product.DoesNotExist:
                pass
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({'success': False})

    # --- PHẦN 1: Quân số ---
    total_workers = WorkerInfo.objects.count()
    attendance_today = AttendanceRecord.objects.filter(date=today)
    present_statuses = ['present', 'presenttg', 'gostation', 'studynuangayinZ', 'studyfullinZ']
    present_count = attendance_today.filter(status__in=present_statuses).count()
    leave_count = attendance_today.filter(status='leave').count()
    study_count = attendance_today.filter(status='study').count()
    pregnancy_count = attendance_today.filter(status='pregnancy').count()
    nghichohuu_count = attendance_today.filter(status='nghichohuu').count()
    curve_count = attendance_today.filter(status='curve').count()
    absent_count = attendance_today.filter(status='absent').count()
    kinhte_count = attendance_today.filter(status='kinhte').count()
    other_count = nghichohuu_count + curve_count + absent_count + kinhte_count
    total_absent = leave_count + study_count + pregnancy_count + nghichohuu_count + curve_count + absent_count + kinhte_count

    # --- PHẦN 2: Tiến độ ---
    product_groups = ProductGroup.objects.all().order_by('name')
    group_data = []
    vector_data = {}

    for pg in product_groups:
        products = Product.objects.filter(item__group=pg)
        latest_progresses = []
        product_detail = {}

        for product in products:
            daily_qs = ProductDailyProgress.objects.filter(product=product).order_by('date')
            daily = {dp.date: {'plan': float(dp.planned_progress_percent), 'real': float(dp.progress_percent),
                               'approved_factory': dp.approved_factory, 'approved_department': dp.approved_department,
                               'approved_qc': dp.approved_qc, 'handed_over': dp.handed_over, 'date': dp.date}
                     for dp in daily_qs}
            if not daily:
                continue
            start_day = min(daily.keys())
            last_data_day = max(daily.keys())
            end_day = last_data_day

            dates, plan, real = [], [], []
            daily_dates = sorted(daily.keys())
            current = start_day
            while current <= end_day:
                if current in daily:
                    dp = daily[current]
                    plan_v, real_v = dp['plan'], dp['real']
                    factory = dp['approved_factory']
                    department = dp['approved_department']
                    qc = dp['approved_qc']
                    handover = dp['handed_over']
                else:
                    last_dates = [d for d in daily_dates if d < current]
                    next_dates = [d for d in daily_dates if d > current]
                    last_date = max(last_dates) if last_dates else None
                    next_date = min(next_dates) if next_dates else None
                    if last_date and next_date:
                        total_days = (next_date - last_date).days
                        elapsed_days = (current - last_date).days
                        plan_v = daily[last_date]['plan'] + (daily[next_date]['plan'] - daily[last_date]['plan']) * elapsed_days / total_days
                        real_v = daily[last_date]['real'] + (daily[next_date]['real'] - daily[last_date]['real']) * elapsed_days / total_days
                    elif last_date:
                        plan_v = daily[last_date]['plan']
                        real_v = daily[last_date]['real']
                    elif next_date:
                        plan_v = daily[next_date]['plan']
                        real_v = daily[next_date]['real']
                    else:
                        plan_v = real_v = 0
                    factory = department = qc = handover = False

                dates.append(current.strftime("%d-%m"))
                plan.append(round(plan_v, 2))
                real.append(round(real_v, 2))
                current += timedelta(days=1)

            completed_steps = ProductStepProgress.objects.filter(product=product, progress_percent__gte=100).order_by('step__order')
            next_step = product.item.steps.exclude(id__in=completed_steps.values('step')).order_by('order').first()
            last_step = completed_steps.last().step if completed_steps.exists() else None

            step_ranges = []
            all_steps = product.item.steps.order_by('order')
            for step in all_steps:
                sp = ProductStepProgress.objects.filter(product=product, step=step).order_by('completed_date').first()
                if sp:
                    start_date = end_date = sp.completed_date
                else:
                    start_date = end_date = None
                if completed_steps.filter(step=step).exists():
                    status = 'done'
                elif step == (next_step if next_step else None):
                    status = 'doing'
                else:
                    status = 'pending'
                if start_date:
                    step_ranges.append({'step_name': step.name, 'status': status,
                                        'start': start_date.strftime("%d-%m"),
                                        'end': end_date.strftime("%d-%m")})

            notes = ProductNote.objects.filter(product=product).select_related('created_by').order_by('-created_at')
            latest_dp = daily[last_data_day]
            product_detail[product.name] = {
                'id': product.id,
                'dates': dates,
                'plan': plan,
                'real': real,
                'approved_factory': latest_dp['date'].strftime("%d/%m/%Y") if latest_dp['approved_factory'] else None,
                'approved_department': latest_dp['date'].strftime("%d/%m/%Y") if latest_dp['approved_department'] else None,
                'approved_qc': latest_dp['date'].strftime("%d/%m/%Y") if latest_dp['approved_qc'] else None,
                'handed_over': latest_dp['date'].strftime("%d/%m/%Y") if latest_dp['handed_over'] else None,
                'current_step': last_step.name if last_step else None,
                'next_step': next_step.name if next_step else None,
                'step_ranges': step_ranges,
                'notes': [{'content': n.content,
                           'author': n.created_by.get_full_name() or n.created_by.username if n.created_by else 'Ẩn danh',
                           'created_at': localtime(n.created_at).strftime("%d/%m/%Y %H:%M")} for n in notes]

            }

            latest = ProductDailyProgress.objects.filter(product=product, date__lte=today).order_by('-date').first()
            if latest:
                latest_progresses.append(latest)

        avg_plan = sum(float(p.planned_progress_percent) for p in latest_progresses) / len(latest_progresses) if latest_progresses else 0
        avg_real = sum(float(p.progress_percent) for p in latest_progresses) / len(latest_progresses) if latest_progresses else 0
        group_data.append({'name': pg.name, 'avg_plan': round(avg_plan, 2), 'avg_real': round(avg_real, 2)})
        vector_data[pg.name] = product_detail

    latest_plan_total = sum(pdata['plan'][-1] for products in vector_data.values() for pdata in products.values() if pdata['plan'])
    latest_real_total = sum(pdata['real'][-1] for products in vector_data.values() for pdata in products.values() if pdata['real'])

    return render(request, 'dashboard_sum.html', {
        'today': today,
        'total_workers': total_workers,
        'present_count': present_count,
        'leave_count': leave_count,
        'study_count': study_count,
        'pregnancy_count': pregnancy_count,
        'other_count': other_count,
        'total_absent': total_absent,
        'group_data': group_data,
        'vector_data': vector_data,
        'latest_plan_total': latest_plan_total,
        'latest_real_total': latest_real_total,
    })



# ---------------------------
# ENTER PRODUCT PROGRESS
# ---------------------------
@login_required
def enter_product_progress(request):
    selected_date = request.GET.get('selected_date')
    if not selected_date:
        selected_date = timezone.localtime(timezone.now()).date()
    else:
        selected_date = timezone.datetime.strptime(selected_date, "%Y-%m-%d").date()

    products = Product.objects.select_related('item__group').order_by('item__group__name', 'name')

    product_info = {}
    for product in products:
        progress = ProductDailyProgress.objects.filter(product=product, date=selected_date).first()
        if not progress:
            progress = ProductDailyProgress.objects.filter(product=product, date__lt=selected_date).order_by('-date').first()
        if not progress:
            progress = ProductDailyProgress(
                product=product,
                progress_percent=0,
                planned_progress_percent=0,
                approved_factory=False,
                approved_department=False,
                approved_qc=False,
                handed_over=False
            )

        completed_steps = ProductStepProgress.objects.filter(product=product, progress_percent__gt=0).order_by('step__order')
        last_step = completed_steps.last().step if completed_steps.exists() else None
        next_step = product.item.steps.filter(order__gt=last_step.order).first() if last_step else product.item.steps.first()
        note = ProductNote.objects.filter(product=product).order_by('-created_at').first()

        product_info[product.id] = {
            'product': product,
            'progress': progress,
            'steps_info': {
                'current_step': last_step,
                'next_step': next_step,
            },
            'note': note,
        }

    product_groups = defaultdict(list)
    for product in products:
        group = product.item.group
        product_groups[group].append(product_info[product.id])

    if request.method == "POST":
        for product in products:
            real_value = request.POST.get(f"real_{product.id}")
            plan_value = request.POST.get(f"plan_{product.id}")
            approved_factory = bool(request.POST.get(f"factory_{product.id}"))
            approved_department = bool(request.POST.get(f"department_{product.id}"))
            approved_qc = bool(request.POST.get(f"qc_{product.id}"))
            handed_over = bool(request.POST.get(f"handover_{product.id}"))

            progress, created = ProductDailyProgress.objects.get_or_create(
                product=product,
                date=selected_date
            )

            if created:
                last_progress = ProductDailyProgress.objects.filter(product=product, date__lt=selected_date).order_by('-date').first()
                if last_progress:
                    progress.progress_percent = last_progress.progress_percent
                    progress.planned_progress_percent = last_progress.planned_progress_percent
                    progress.approved_factory = last_progress.approved_factory
                    progress.approved_department = last_progress.approved_department
                    progress.approved_qc = last_progress.approved_qc
                    progress.handed_over = last_progress.handed_over

            if real_value:
                progress.progress_percent = float(real_value)
            if plan_value:
                progress.planned_progress_percent = float(plan_value)
            progress.approved_factory = approved_factory
            progress.approved_department = approved_department
            progress.approved_qc = approved_qc
            progress.handed_over = handed_over

            progress.save()

        messages.success(request, "Cập nhật tiến độ thành công!")
        return redirect(f"{request.path}?selected_date={selected_date}")

    return render(request, 'enter_product_progress.html', {
        'selected_date': selected_date,
        'product_groups': dict(product_groups),
    })
